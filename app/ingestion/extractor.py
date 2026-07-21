"""
Text extraction for the RAG ingestion pipeline.
Handles: digital PDFs, scanned/image PDFs (via OCR), and DOCX files.

Dependencies:
    pip install pymupdf pytesseract pillow python-docx --break-system-packages
    # Tesseract binary must also be installed on the system (apt install tesseract-ocr)
"""

import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
from docx import Document
from dataclasses import dataclass
from pathlib import Path


@dataclass
class PageContent:
    page_number: int
    text: str
    source_id: str
    was_ocr: bool = False


# Below this char count per page, assume the page is scanned/image-only
OCR_FALLBACK_THRESHOLD = 50


def extract_pdf(file_path: str, source_id: str) -> list[PageContent]:
    """Extract text from a PDF, falling back to OCR per-page when needed."""
    doc = fitz.open(file_path)
    pages = []

    for i, page in enumerate(doc):
        text = page.get_text().strip()

        if len(text) < OCR_FALLBACK_THRESHOLD:
            # Likely a scanned page — render to image and OCR it
            pix = page.get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text = pytesseract.image_to_string(img).strip()
            pages.append(PageContent(page_number=i + 1, text=text, source_id=source_id, was_ocr=True))
        else:
            pages.append(PageContent(page_number=i + 1, text=text, source_id=source_id, was_ocr=False))

    doc.close()
    return pages


def extract_docx(file_path: str, source_id: str) -> list[PageContent]:
    """Extract text from a DOCX file. DOCX has no native page concept,
    so we treat the whole document as one logical unit (page_number=1);
    downstream chunking will split it regardless."""
    doc = Document(file_path)

    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # Pull table content too — often holds key data (specs, pricing, etc.)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    full_text = "\n".join(parts)
    return [PageContent(page_number=1, text=full_text, source_id=source_id, was_ocr=False)]


def extract(file_path: str, source_id: str) -> list[PageContent]:
    """Dispatch to the right extractor based on file extension."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return extract_pdf(file_path, source_id)
    elif ext == ".docx":
        return extract_docx(file_path, source_id)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


if __name__ == "__main__":
    # Quick manual test
    import sys
    if len(sys.argv) > 1:
        result = extract(sys.argv[1], source_id="test-doc")
        for page in result:
            print(f"--- Page {page.page_number} (OCR: {page.was_ocr}) ---")
            print(page.text[:300])
            print()
